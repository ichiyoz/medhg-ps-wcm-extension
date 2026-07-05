"""Export Tables 2 and 3 (with this study's new rows) to a single .docx."""
from pathlib import Path

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path.home() / "Dropbox" / "Surgery" / "Tables_2_3.docx"

T2_CAP = ("Table 2. Performance of the deployable models and candidate under five-fold "
          "cross-validation (n = 44,721; event rate 9.6%; no-skill AUPRC = 0.096). "
          "Point estimates are the cross-validation mean; parentheses give the 95% "
          "confidence interval across the five folds. Probabilities are isotonic-"
          "calibrated. The gradient-boosted tree with the GRU care-path sequence "
          "encoder is a deployable candidate added in this study.")
T2_HEAD = ["Model", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier"]
T2_ROWS = [
    ["Gradient-boosted tree + GRU care-path sequence encoder (clinical features and learned care path)",
     "0.703 (0.699–0.707)", "0.242 (0.229–0.255)", "0.081"],
    ["Gradient-boosted tree (clinical features and care path sequences)",
     "0.703 (0.696–0.710)", "0.235 (0.218–0.252)", "0.082"],
    ["Gradient-boosted tree (clinical features only)",
     "0.703 (0.692–0.714)", "0.225 (0.207–0.243)", "0.082"],
    ["MedHG-PS extension: Graph-neural-network embedding + Gradient-boosted tree "
     "(clinical features and care path sequences)",
     "0.700 (0.692–0.708)", "0.233 (0.218–0.248)", "0.081"],
]

T3_CAP = ("Table 3. Graph-based models tested (five-fold cross-validation; n = 44,721). "
          "Each row is the best configuration of its family (tuned per fold). Point "
          "estimates are the cross-validation mean; parentheses give the 95% confidence "
          "interval across the five folds (no-skill AUPRC = 0.096). None matched the "
          "best models in Table 2.")
T3_HEAD = ["Graph-based model", "AUROC (95% CI)", "AUPRC (95% CI)"]
# (label, auroc, auprc, is_section_header)
T3_ROWS = [
    ("Graph-neural-network embedding integration", "", "", True),
    ("Original MedHG-PS", "0.685 (0.673–0.697)", "0.193 (0.174–0.212)", False),
    ("Graph-neural-network embedding (encounter) + tree", "0.685 (0.671–0.699)", "0.207 (0.186–0.228)", False),
    ("Graph-neural-network embedding (provider/unit) + tree", "0.694 (0.683–0.705)", "0.226 (0.211–0.241)", False),
    ("Graph-neural-network embedding (all node types) + tree", "0.687 (0.676–0.698)", "0.210 (0.196–0.224)", False),
    ("Higher-order and standard architectures", "", "", True),
    ("Hypergraph network, mean aggregation", "0.688 (0.677–0.699)", "0.210 (0.202–0.218)", False),
    ("Hypergraph network, hub-pruned", "0.683 (0.668–0.698)", "0.206 (0.194–0.218)", False),
    ("Hypergraph network, attention aggregation", "0.690 (0.672–0.708)", "0.212 (0.194–0.230)", False),
    ("Conjunctive hyperedges (best of >20)", "0.693 (0.676–0.710)", "0.209 (0.192–0.226)", False),
    ("Standard graph-neural-network architectures (best of four)", "0.692 (0.684–0.700)", "0.201 (0.189–0.213)", False),
    ("Clinical-concept-node heterograph (diagnosis + procedure nodes)", "0.692 (0.682–0.702)", "0.198 (0.184–0.212)", False),
    ("Sequence and similarity extensions (this study)", "", "", True),
    ("GRU care-path sequence encoder (sequence only)", "0.665 (0.658–0.672)", "0.186 (0.179–0.193)", False),
    ("Patient-similarity-edge GNN (encounter–encounter kNN)", "0.649 (0.634–0.664)", "0.168 (0.150–0.186)", False),
    ("Temporal-edge GNN (edge-conditioned MPNN)", "0.695 (0.685–0.705)", "0.202 (0.187–0.217)", False),
]


def _bold_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def main():
    doc = docx.Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(54)

    cap = doc.add_paragraph(); cap.add_run(T2_CAP).bold = False
    cap.paragraph_format.space_after = Pt(6)
    t2 = doc.add_table(rows=1, cols=len(T2_HEAD)); t2.style = "Table Grid"
    for c, h in zip(t2.rows[0].cells, T2_HEAD):
        c.paragraphs[0].add_run(h)
    _bold_row(t2.rows[0])
    for r in T2_ROWS:
        cells = t2.add_row().cells
        for c, v in zip(cells, r):
            c.paragraphs[0].add_run(v)

    doc.add_paragraph()
    cap = doc.add_paragraph(); cap.add_run(T3_CAP).bold = False
    cap.paragraph_format.space_after = Pt(6)
    t3 = doc.add_table(rows=1, cols=len(T3_HEAD)); t3.style = "Table Grid"
    for c, h in zip(t3.rows[0].cells, T3_HEAD):
        c.paragraphs[0].add_run(h)
    _bold_row(t3.rows[0])
    for label, auroc, auprc, is_hdr in T3_ROWS:
        cells = t3.add_row().cells
        if is_hdr:
            a = cells[0].merge(cells[1]).merge(cells[2])
            run = a.paragraphs[0].add_run(label); run.bold = True; run.italic = True
        else:
            cells[0].paragraphs[0].add_run(label)
            for ci, val in ((1, auroc), (2, auprc)):
                p = cells[ci].paragraphs[0]; p.add_run(val)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"wrote {OUT}  (Table 2: {len(T2_ROWS)} rows, Table 3: {len(T3_ROWS)} rows incl. headers)")


if __name__ == "__main__":
    main()
